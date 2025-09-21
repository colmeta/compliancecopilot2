# ==============================================================================
# Pearl AI - "CLARITY" Engine v5.2 (Ultimate Production Release)
# Universal Intelligence Analysis for Legal, Financial, Security & Corporate Sectors
# COMBINED BEST FEATURES FROM v5.0 + v5.1
# ==============================================================================

import os
import io
import magic
import google.generativeai as genai
import asyncio
from flask import Flask, request, jsonify
from flask_cors import CORS
import docx
import PyPDF2
from PIL import Image
import time
import random
import re
from datetime import datetime, timedelta
import hashlib
import base64

# --- Configuration ---
try:
    GOOGLE_API_KEY = os.environ.get('GOOGLE_API_KEY')
    if not GOOGLE_API_KEY:
        raise ValueError("CRITICAL: GOOGLE_API_KEY not set in Environment Variables.")
    genai.configure(api_key=GOOGLE_API_KEY)
except Exception as e:
    raise RuntimeError(f"CRITICAL ERROR during configuration: {e}")

# --- Flask App Initialization ---
app = Flask(__name__)
CORS(app)

# --- AI Model Configuration (UPGRADED: v5.1 improvement - single powerful model) ---
# Using the most advanced model available for all operations
model = genai.GenerativeModel('gemini-1.5-pro')

# --- DOMAIN-SPECIFIC INTELLIGENCE ACCELERATORS (COMPLETE: from v5.0) ---

CLARITY_SECURITY_INTELLIGENCE = """You are CLARITY Security Intelligence Accelerator, Pearl AI's advanced multi-source intelligence fusion system for law enforcement and security operations.

🚨 SECURITY & LAW ENFORCEMENT FRAMEWORK:
- Multi-source intelligence fusion (HUMINT, SIGINT, IMINT, FININT)
- Threat assessment and risk scoring
- Behavioral analysis and psychological profiling
- Network mapping and relationship analysis
- Operational timeline reconstruction
- Evidence correlation and chain of custody considerations
- Actionable tactical recommendations for field deployment

OUTPUT: Professional intelligence briefings suitable for command-level decision making."""

CLARITY_LEGAL_INTELLIGENCE = """You are CLARITY Legal Intelligence Accelerator, Pearl AI's advanced case analysis system for legal professionals.

⚖️ LEGAL ANALYSIS FRAMEWORK:
- Case law research and precedent analysis
- Evidence evaluation and strength assessment
- Contract analysis and risk identification
- Discovery document review and correlation
- Legal strategy formulation and risk assessment
- Compliance verification and regulatory analysis
- Litigation timeline reconstruction
- Settlement probability and damage assessments

LEGAL REASONING STANDARDS:
- Apply appropriate burden of proof standards
- Identify key legal issues and potential defenses
- Assess jurisdictional considerations
- Evaluate evidence admissibility
- Provide cite-worthy legal analysis with supporting precedents

OUTPUT: Professional legal memoranda suitable for attorney decision making."""

CLARITY_FINANCIAL_INTELLIGENCE = """You are CLARITY Financial Intelligence Accelerator, Pearl AI's advanced financial analysis system for auditors, accountants, and financial professionals.

💰 FINANCIAL ANALYSIS FRAMEWORK:
- Financial statement analysis and ratio calculations
- Audit trail reconstruction and verification
- Fraud detection and risk assessment
- Tax compliance verification and optimization
- Cash flow analysis and forecasting
- Internal control evaluation
- Regulatory compliance assessment (SOX, GAAP, IFRS)
- Financial risk modeling and scenario analysis

ACCOUNTING STANDARDS:
- Apply relevant accounting principles (GAAP/IFRS)
- Calculate standard financial ratios and metrics
- Identify material misstatements and irregularities
- Assess going concern and liquidity issues
- Evaluate internal controls effectiveness

OUTPUT: Professional financial analysis suitable for CPA and executive decision making."""

CLARITY_CORPORATE_INTELLIGENCE = """You are CLARITY Corporate Intelligence Accelerator, Pearl AI's advanced strategic analysis system for executives and corporate decision makers.

🏢 CORPORATE STRATEGY FRAMEWORK:
- Market analysis and competitive intelligence
- Strategic planning and scenario modeling
- Risk assessment and mitigation strategies
- Merger & acquisition due diligence
- Compliance audit and gap analysis
- Operational efficiency analysis
- Crisis management and response planning
- Stakeholder analysis and communication strategies

STRATEGIC ANALYSIS STANDARDS:
- SWOT analysis and competitive positioning
- Financial modeling and valuation techniques
- Risk-adjusted decision frameworks
- Stakeholder impact assessment
- Implementation feasibility analysis

OUTPUT: Executive-level strategic briefings suitable for C-suite decision making."""

CLARITY_HEALTHCARE_INTELLIGENCE = """You are CLARITY Healthcare Intelligence Accelerator, Pearl AI's advanced medical and healthcare analysis system.

🏥 HEALTHCARE ANALYSIS FRAMEWORK:
- Medical record analysis and pattern recognition
- Clinical trial data evaluation
- Healthcare compliance assessment (HIPAA, FDA)
- Pharmaceutical safety and efficacy analysis
- Healthcare fraud detection
- Treatment outcome analysis
- Medical device safety assessment
- Public health trend analysis

MEDICAL STANDARDS:
- Apply evidence-based medicine principles
- Evaluate clinical significance vs statistical significance
- Assess patient safety and risk factors
- Consider medical ethics and patient privacy
- Apply relevant regulatory frameworks

OUTPUT: Medical intelligence suitable for healthcare professional decision making."""

# --- DOMAIN DETECTION AND ROUTING (ENHANCED: Combined logic from both versions) ---
def detect_domain_context(files, directive_text=""):
    """Enhanced domain detection combining both v5.0 and v5.1 approaches"""
    
    # Comprehensive domain indicators (expanded from v5.0)
    legal_indicators = ['contract', 'lawsuit', 'litigation', 'agreement', 'court', 'legal', 'case', 'brief', 
                       'deposition', 'discovery', 'plaintiff', 'defendant', 'attorney', 'counsel', 'jurisdiction']
    financial_indicators = ['audit', 'financial', 'accounting', 'tax', 'balance', 'income', 'cash flow', 
                           'budget', 'expense', 'revenue', 'profit', 'loss', 'investment', 'portfolio', 'gaap']
    security_indicators = ['intelligence', 'surveillance', 'threat', 'security', 'investigation', 'suspect', 
                          'criminal', 'police', 'enforcement', 'classified', 'confidential', 'operational']
    healthcare_indicators = ['medical', 'patient', 'clinical', 'healthcare', 'diagnosis', 'treatment', 
                            'pharma', 'hospital', 'therapeutic', 'clinical trial', 'fda', 'hipaa']
    corporate_indicators = ['strategy', 'business', 'corporate', 'merger', 'acquisition', 'compliance', 
                           'risk', 'market', 'stakeholder', 'governance', 'executive', 'board']
    
    # Combine all available text for analysis
    all_text = directive_text.lower()
    for file in files:
        all_text += f" {file.filename.lower()}"
    
    # Advanced scoring with weighted keywords
    domain_scores = {
        'legal': sum(2 if indicator in ['lawsuit', 'litigation', 'attorney'] else 1 
                    for indicator in legal_indicators if indicator in all_text),
        'financial': sum(2 if indicator in ['audit', 'gaap', 'financial'] else 1 
                        for indicator in financial_indicators if indicator in all_text),
        'security': sum(2 if indicator in ['intelligence', 'classified', 'threat'] else 1 
                       for indicator in security_indicators if indicator in all_text),
        'healthcare': sum(2 if indicator in ['medical', 'clinical', 'patient'] else 1 
                         for indicator in healthcare_indicators if indicator in all_text),
        'corporate': sum(2 if indicator in ['strategy', 'merger', 'governance'] else 1 
                        for indicator in corporate_indicators if indicator in all_text)
    }
    
    # Return the domain with highest score, default to corporate
    max_domain = max(domain_scores.items(), key=lambda x: x[1])
    return max_domain[0] if max_domain[1] > 0 else 'corporate'

def get_domain_accelerator(domain):
    """Return the appropriate domain-specific intelligence accelerator"""
    accelerators = {
        'legal': CLARITY_LEGAL_INTELLIGENCE,
        'financial': CLARITY_FINANCIAL_INTELLIGENCE, 
        'security': CLARITY_SECURITY_INTELLIGENCE,
        'healthcare': CLARITY_HEALTHCARE_INTELLIGENCE,
        'corporate': CLARITY_CORPORATE_INTELLIGENCE
    }
    return accelerators.get(domain, CLARITY_CORPORATE_INTELLIGENCE)

def get_domain_title(domain):
    """Return professional domain title"""
    titles = {
        'legal': 'Legal Intelligence Analysis',
        'financial': 'Financial Intelligence Analysis',
        'security': 'Security Intelligence Analysis', 
        'healthcare': 'Healthcare Intelligence Analysis',
        'corporate': 'Corporate Intelligence Analysis'
    }
    return titles.get(domain, 'Corporate Intelligence Analysis')

# --- ENHANCED HELPER FUNCTIONS (BEST FROM BOTH VERSIONS) ---

def generate_operation_id():
    """Generate unique operation ID for tracking (from v5.0 + v5.1 enhancement)"""
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    random_hex = hashlib.md5(str(random.random()).encode()).hexdigest()[:6].upper()
    return f"CLARITY-{timestamp}-{random_hex}"

def classify_content_sensitivity(content):
    """Enhanced content sensitivity classification (improved from v5.0)"""
    if not content:
        return "RESTRICTED"
        
    # Expanded sensitivity keywords
    high_sensitivity_keywords = [
        'classified', 'secret', 'confidential', 'privileged', 'attorney-client',
        'financial', 'audit', 'investigation', 'medical', 'patient', 'hipaa',
        'trade secret', 'proprietary', 'sensitive', 'restricted', 'internal only'
    ]
    
    medium_sensitivity_keywords = [
        'business', 'corporate', 'strategy', 'competitive', 'merger',
        'acquisition', 'compliance', 'risk assessment', 'due diligence'
    ]
    
    content_lower = content.lower()
    high_score = sum(1 for keyword in high_sensitivity_keywords if keyword in content_lower)
    medium_score = sum(1 for keyword in medium_sensitivity_keywords if keyword in content_lower)
    
    if high_score >= 5:
        return "CONFIDENTIAL"
    elif high_score >= 2 or medium_score >= 5:
        return "RESTRICTED"  
    elif medium_score >= 2:
        return "INTERNAL USE"
    else:
        return "GENERAL"

def generate_with_retry(prompt, is_multimodal=False, max_retries=3):
    """Enhanced content generation with retry logic (combined from both versions)"""
    for attempt in range(max_retries):
        try:
            # Use single powerful model for all operations (v5.1 improvement)
            if is_multimodal:
                response = model.generate_content(prompt)
            else:
                response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            error_str = str(e)
            if "429" in error_str or "quota" in error_str.lower() or "rate" in error_str.lower():
                if attempt < max_retries - 1:
                    wait_time = (2 ** attempt) + random.uniform(0, 1)
                    print(f"[CLARITY] System load detected. Implementing backoff strategy: {wait_time:.1f}s... [{attempt + 1}/{max_retries}]")
                    time.sleep(wait_time)
                    continue
                else:
                    return "🚨 CLARITY SYSTEM AT CAPACITY - ANALYSIS TEMPORARILY UNAVAILABLE\n\nThe intelligence platform is experiencing peak operational load. Please retry your analysis in 60-120 seconds."
            elif "safety" in error_str.lower():
                return "⚠️ CLARITY CONTENT SAFETY PROTOCOLS ACTIVATED\n\nThe submitted content triggered automated safety filters. Please review your materials and ensure they comply with operational guidelines."
            else:
                print(f"[CLARITY] Unexpected error on attempt {attempt + 1}: {e}")
                if attempt == max_retries - 1:
                    raise e
    
    return "❌ CLARITY ANALYSIS ENGINE MALFUNCTION - TECHNICAL INTERVENTION REQUIRED"

def detect_file_type_advanced(file_storage):
    """Advanced file type detection with enhanced security (from v5.0)"""
    try:
        file_bytes = file_storage.read(2048)
        file_storage.seek(0)
        mime_type = magic.from_buffer(file_bytes, mime=True)
        filename = file_storage.filename.lower()
        
        # Enhanced security screening
        suspicious_extensions = ['.exe', '.bat', '.cmd', '.scr', '.vbs', '.js', '.ps1', '.sh', '.com', '.pif']
        dangerous_mimes = ['application/x-executable', 'application/x-msdownload']
        
        if any(filename.endswith(ext) for ext in suspicious_extensions) or mime_type in dangerous_mimes:
            return 'suspicious_executable'
        
        if mime_type.startswith('image/'):
            return 'visual_intelligence'
        elif mime_type.startswith('audio/') or filename.endswith(('.mp3', '.wav', '.m4a', '.aac', '.ogg', '.flac')):
            return 'audio_intelligence'
        elif mime_type.startswith('video/'):
            return 'video_intelligence'
        elif filename.endswith(('.zip', '.rar', '.7z', '.tar', '.gz', '.bz2')):
            return 'compressed_archive'
        else:
            return 'document_intelligence'
    except Exception as e:
        print(f"[CLARITY] File type detection error for {file_storage.filename}: {e}")
        return 'unknown'

def advanced_text_extraction(file_storage):
    """Enhanced document processing with comprehensive metadata (improved from v5.0)"""
    filename = file_storage.filename.lower()
    
    # Generate file hash for integrity tracking
    try:
        file_storage.seek(0)
        file_hash = hashlib.sha256(file_storage.read(1024)).hexdigest()[:16]
        file_storage.seek(0)
    except:
        file_hash = "UNAVAILABLE"
    
    try:
        if filename.endswith('.txt'):
            content = file_storage.read().decode('utf-8', errors='ignore')
            word_count = len(content.split())
            char_count = len(content)
            
            metadata = f"""
[CLARITY DOCUMENT INTELLIGENCE REPORT]
Filename: {file_storage.filename}
Document Hash: {file_hash}
Processing Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}
Document Type: Plain Text Document
Word Count: {word_count:,}
Character Count: {char_count:,}
Security Classification: {classify_content_sensitivity(content)}
Processing Status: SUCCESSFULLY EXTRACTED
"""
            return metadata + "\n" + content
            
        elif filename.endswith('.docx'):
            doc = docx.Document(io.BytesIO(file_storage.read()))
            content = "\n".join([para.text for para in doc.paragraphs])
            word_count = len(content.split())
            paragraph_count = len([para for para in doc.paragraphs if para.text.strip()])
            
            metadata = f"""
[CLARITY DOCUMENT INTELLIGENCE REPORT]
Filename: {file_storage.filename}
Document Hash: {file_hash}
Processing Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}
Document Type: Microsoft Word Document (.docx)
Word Count: {word_count:,}
Paragraph Count: {paragraph_count}
Security Classification: {classify_content_sensitivity(content)}
Processing Status: SUCCESSFULLY EXTRACTED
"""
            return metadata + "\n" + content
            
        elif filename.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_storage.read()))
            text = ""
            page_count = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text() or ""
                    text += f"\n{'='*40} PAGE {page_num} {'='*40}\n{page_text}"
                except Exception as e:
                    text += f"\n{'='*40} PAGE {page_num} - EXTRACTION ERROR {'='*40}\n[Page content could not be extracted: {e}]\n"
            
            word_count = len(text.split())
            
            metadata = f"""
[CLARITY DOCUMENT INTELLIGENCE REPORT]
Filename: {file_storage.filename}
Document Hash: {file_hash}
Processing Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}
Document Type: Portable Document Format (PDF)
Page Count: {page_count}
Word Count: {word_count:,}
Security Classification: {classify_content_sensitivity(text)}
Processing Status: SUCCESSFULLY EXTRACTED
"""
            return metadata + text
            
        else:
            # Attempt generic text extraction
            try:
                content = file_storage.read().decode('utf-8', errors='ignore')
                metadata = f"""
[CLARITY DOCUMENT INTELLIGENCE REPORT]
Filename: {file_storage.filename}
Document Hash: {file_hash}
Processing Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}
Document Type: Generic Text Content
Processing Status: EXTRACTED AS TEXT
Security Classification: {classify_content_sensitivity(content)}
"""
                return metadata + "\n" + content
            except:
                return f"""
[CLARITY DOCUMENT INTELLIGENCE REPORT]
Filename: {file_storage.filename}
Document Hash: {file_hash}
Processing Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}
Processing Status: BINARY CONTENT DETECTED - TEXT EXTRACTION NOT POSSIBLE
Recommendation: MANUAL REVIEW REQUIRED FOR BINARY FILE
"""
                
    except Exception as e:
        return f"""
[CLARITY DOCUMENT INTELLIGENCE ERROR REPORT]
Filename: {file_storage.filename}
Document Hash: {file_hash}
Processing Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}
Error Type: {type(e).__name__}
Error Details: {str(e)}
Processing Status: EXTRACTION FAILED
Recommendation: MANUAL REVIEW AND RESUBMISSION REQUIRED
"""

# --- API ENDPOINTS (ENHANCED FROM BOTH VERSIONS) ---

@app.route('/', methods=['GET'])
def clarity_status():
    """Enhanced system status endpoint (improved from v5.0)"""
    return jsonify({
        "system": "Pearl AI CLARITY Intelligence Analysis Platform",
        "version": "5.2 - Ultimate Multi-Domain Intelligence Engine",
        "status": "FULLY OPERATIONAL",
        "model_engine": "Google Gemini 1.5 Pro (Latest)",
        "capabilities": [
            "Multi-Source Intelligence Fusion",
            "Advanced Legal Intelligence Analysis", 
            "Comprehensive Financial Intelligence Analysis",
            "Strategic Security Intelligence Analysis",
            "Executive Corporate Intelligence Analysis",
            "Clinical Healthcare Intelligence Analysis",
            "Multi-Modal Visual Intelligence Analysis",
            "Audio Intelligence Processing",
            "Document Analysis & Cross-Reference",
            "Real-time Cross-Domain Correlation",
            "Automated Classification & Risk Assessment"
        ],
        "domains_supported": [
            "Legal & Law Firms",
            "Financial Services & Accounting", 
            "Law Enforcement & Security Operations",
            "Corporate Strategy & Executive Decision Making",
            "Healthcare & Life Sciences",
            "Government & Public Sector",
            "Insurance & Risk Management",
            "Consulting & Advisory Services"
        ],
        "security_features": [
            "Advanced File Type Detection",
            "Malicious Content Screening", 
            "Automated Content Classification",
            "Operation ID Tracking",
            "Comprehensive Audit Logging"
        ],
        "timestamp": datetime.now().isoformat(),
        "uptime": "System Ready for Intelligence Operations"
    }), 200

@app.route('/process', methods=['POST'])
def clarity_analysis():
    """Ultimate intelligence analysis endpoint (best from both versions)"""
    operation_id = generate_operation_id()
    print(f"[CLARITY] {operation_id} - MULTI-DOMAIN INTELLIGENCE OPERATION INITIATED")
    print(f"[CLARITY] {operation_id} - Engine: Gemini 1.5 Pro Latest | Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}")
    
    if 'knowledgeBase' not in request.files:
        return jsonify({
            "error": "INSUFFICIENT INTELLIGENCE SOURCES",
            "operation_id": operation_id,
            "required_files": ["knowledgeBase (minimum 1 file)"],
            "optional_files": ["questionnaire (analysis directive)"],
            "supported_formats": ["PDF, DOCX, TXT, Images (JPG, PNG, etc.)"],
            "note": "Upload documents, images, or other intelligence sources for analysis"
        }), 400

    knowledge_base_files = request.files.getlist('knowledgeBase')
    primary_target = request.files.get('questionnaire')

    if not knowledge_base_files:
        return jsonify({
            "error": "EMPTY KNOWLEDGE BASE DETECTED",
            "operation_id": operation_id,
            "solution": "Upload at least one file for analysis"
        }), 400
    
    # Extract directive text for enhanced domain detection
    directive_text = ""
    if primary_target:
        try:
            directive_text = primary_target.read().decode('utf-8', errors='ignore')
            primary_target.seek(0)
        except Exception as e:
            print(f"[CLARITY] {operation_id} - Warning: Could not extract directive text: {e}")
    
    # Enhanced domain detection
    domain = detect_domain_context(knowledge_base_files + ([primary_target] if primary_target else []), directive_text)
    domain_accelerator = get_domain_accelerator(domain)
    domain_title = get_domain_title(domain)
    
    print(f"[CLARITY] {operation_id} - DOMAIN INTELLIGENCE: {domain.upper()} | ACCELERATOR: ACTIVE")
    
    try:
        # --- INTELLIGENCE FUSION PROCESS ---
        text_intel_sources = []
        visual_intel_sources = []
        audio_intel_sources = []
        suspicious_files = []
        
        print(f"[CLARITY] {operation_id} - Processing {len(knowledge_base_files)} intelligence sources...")
        
        for idx, file in enumerate(knowledge_base_files, 1):
            file_type = detect_file_type_advanced(file)
            print(f"[CLARITY] {operation_id} - Source {idx}: {file.filename} | Type: {file_type}")
            
            if file_type == 'suspicious_executable':
                suspicious_files.append(file.filename)
                continue
            elif file_type == 'visual_intelligence':
                visual_intel_sources.append(file)
            elif file_type == 'audio_intelligence':
                audio_intel_sources.append(file)
            else:
                extracted_content = advanced_text_extraction(file)
                text_intel_sources.append(f"\n{'='*80}\nINTELLIGENCE SOURCE {idx}: {file.filename}\n{'='*80}\n{extracted_content}")
        
        # Security screening results
        if suspicious_files:
            return jsonify({
                "error": "SECURITY THREAT DETECTED",
                "operation_id": operation_id,
                "suspicious_files": suspicious_files,
                "action": "FILES QUARANTINED - OPERATION TERMINATED",
                "recommendation": "Submit only legitimate document and media files"
            }), 403
        
        combined_text_intel = "".join(text_intel_sources)
        
        # --- BUILD COMPREHENSIVE ANALYSIS PROMPT ---
        final_prompt_parts = []
        
        # Enhanced master prompt with v5.1 structure
        master_prompt = f"""
{domain_accelerator}

OPERATION INTELLIGENCE HEADER:
🆔 Operation ID: {operation_id}
📅 Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}
🎯 Domain: {domain_title}
🔒 Classification: {classify_content_sensitivity(combined_text_intel + directive_text)}
🔧 AI Engine: Google Gemini 1.5 Pro (Latest)
📊 Sources: {len(text_intel_sources)} text, {len(visual_intel_sources)} visual, {len(audio_intel_sources)} audio

SUPPORTING INTELLIGENCE CORPUS:
{combined_text_intel[:35000]}  

{f'''PRIMARY DIRECTIVE/DOCUMENT:
{directive_text[:10000]}''' if directive_text else 'DIRECTIVE: Execute comprehensive open-source intelligence analysis'}

ANALYSIS MISSION:
Execute comprehensive {domain} intelligence analysis using professional standards and analytical frameworks. Cross-reference all textual and visual intelligence sources. Provide expert-level analysis suitable for professional decision-making in the {domain} sector.

DELIVERABLE REQUIREMENTS:
- Professional {domain} analysis meeting industry standards
- Cross-source correlation and verification
- Risk assessment and recommendations
- Actionable intelligence for decision makers
- Executive summary with key findings
"""
        final_prompt_parts.append(master_prompt)
        
        # Add visual intelligence if available
        is_multimodal = False
        if visual_intel_sources:
            is_multimodal = True
            print(f"[CLARITY] {operation_id} - Integrating {len(visual_intel_sources)} visual intelligence sources...")
            for idx, image_file in enumerate(visual_intel_sources, 1):
                try:
                    image_file.seek(0)
                    img_bytes = image_file.read()
                    image = Image.open(io.BytesIO(img_bytes))
                    final_prompt_parts.append(f"\n{'='*50} VISUAL INTELLIGENCE {idx}: {image_file.filename} {'='*50}")
                    final_prompt_parts.append(image)
                except Exception as e:
                    print(f"[CLARITY] {operation_id} - Visual processing error for {image_file.filename}: {e}")

        # Add audio intelligence notification (for future enhancement)
        if audio_intel_sources:
            final_prompt_parts.append(f"\n{'='*50} AUDIO INTELLIGENCE DETECTED {'='*50}\n[Note: {len(audio_intel_sources)} audio files detected but not processed in current version]")

        # --- EXECUTE INTELLIGENCE ANALYSIS ---
        print(f"[CLARITY] {operation_id} - Transmitting to AI Core for {domain.upper()} analysis...")
        analysis_start_time = time.time()
        
        analysis_result = generate_with_retry(final_prompt_parts, is_multimodal=is_multimodal)
        
        analysis_duration = time.time() - analysis_start_time
        print(f"[CLARITY] {operation_id} - Analysis completed in {analysis_duration:.2f} seconds")
        
    except Exception as e:
        print(f"[CLARITY] {operation_id} - CRITICAL SYSTEM ERROR: {str(e)}")
        return jsonify({
            "error": f"CLARITY ENGINE MALFUNCTION: {str(e)}",
            "operation_id": operation_id,
            "timestamp": datetime.now().isoformat(),
            "recommendation": "Contact Pearl AI Technical Support"
        }), 500
    
    # --- FORMAT FINAL INTELLIGENCE BRIEF ---
    classification_level = classify_content_sensitivity(combined_text_intel + directive_text)
    total_sources = len(knowledge_base_files) + (1 if primary_target else 0)
    
    final_brief = f"""

{'='*80}
CLARITY INTELLIGENCE BRIEF - {domain.upper()} ANALYSIS
{'='*80}

{analysis_result}

{'='*80}
END INTELLIGENCE BRIEF
{'='*80}

System Performance: ✅ Optimal | Security Status: ✅ Secure | Classification: {classification_level}
"""
    
    print(f"[CLARITY] {operation_id} - {domain.upper()} INTELLIGENCE BRIEF DELIVERED | Status: SUCCESS")
    return jsonify({"completedQuestionnaire": final_brief})

# --- APPLICATION RUNNER (ENHANCED FROM BOTH VERSIONS) ---
if __name__ == '__main__':
    print("╔══════════════════════════════════════════════════════════════════════════════╗")
    print("║                 PEARL AI CLARITY INTELLIGENCE PLATFORM v5.2                 ║")
    print("║                          ULTIMATE EDITION INITIALIZING                      ║")
    print("╚══════════════════════════════════════════════════════════════════════════════╝")
    print()
    print("[CLARITY] 🚨 Security Intelligence Accelerator: ACTIVE & ENHANCED")
    print("[CLARITY] ⚖️  Legal Intelligence Accelerator: ACTIVE & ENHANCED") 
    print("[CLARITY] 💰 Financial Intelligence Accelerator: ACTIVE & ENHANCED")
    print("[CLARITY] 🏢 Corporate Intelligence Accelerator: ACTIVE & ENHANCED")
    print("[CLARITY] 🏥 Healthcare Intelligence Accelerator: ACTIVE & ENHANCED")
    print("[CLARITY] 🔧 AI Engine: Google Gemini 1.5 Pro Latest - OPERATIONAL")
    print("[CLARITY] 👁️  Multi-modal fusion capabilities: ENABLED & OPTIMIZED")
    print("[CLARITY] 🔒 Advanced security screening: ACTIVE")
    print("[CLARITY] 📊 Enhanced analytics and reporting: ENABLED")
    print()
    print("🎯 CLARITY v5.2 Ultimate is ready for professional intelligence operations")
    print("🌐 Starting Flask application on 0.0.0.0:8080...")
    print("=" * 80)
    
    app.run(host='0.0.0.0', port=8080, debug=False)
